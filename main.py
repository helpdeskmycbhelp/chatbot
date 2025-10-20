import os, io, re, json, time, pathlib
from typing import List, Dict, Optional, Literal
from fastapi import FastAPI, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from dotenv import load_dotenv
import numpy as np
import requests

# ----------------------- Config -----------------------
load_dotenv()

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
OPENAI_BASE    = os.getenv("OPENAI_BASE", "https://api.openai.com/v1")
CHAT_MODEL     = os.getenv("MODEL", "gpt-4o-mini")
EMBED_MODEL    = os.getenv("EMBED_MODEL", "text-embedding-3-small")
TOP_K          = int(os.getenv("TOP_K", "4"))
CHUNK_SIZE     = int(os.getenv("CHUNK_SIZE", "1200"))

if not OPENAI_API_KEY:
    raise RuntimeError("Missing OPENAI_API_KEY. Put it in .env")

DATA_DIR    = pathlib.Path("./data"); DATA_DIR.mkdir(exist_ok=True)
EMB_NPY     = DATA_DIR / "embeddings.npy"
CHUNKS_JSON = DATA_DIR / "chunks.json"
META_JSON   = DATA_DIR / "meta.json"

# ----------------------- App & CORS -----------------------
app = FastAPI(title="Policy Chatbot API", version="1.4")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],          # tighten in production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

from fastapi.staticfiles import StaticFiles
app.mount("/static", StaticFiles(directory="static"), name="static")

# ----------------------- Tiny Vector Store -----------------------
KB: Dict[str, Optional[object]] = {"chunks": [], "embeds": None, "meta": []}

def _clean_text(s: str) -> str:
    s = s.replace("\r", "")
    s = re.sub(r"\n{3,}", "\n\n", s)
    return s.strip()

def chunk_text(txt: str, max_chars: int) -> List[str]:
    lines = [l.strip() for l in txt.splitlines()]
    parts, buf = [], []
    for ln in lines:
        candidate = ("\n".join(buf + [ln])).strip()
        if len(candidate) > max_chars and buf:
            parts.append("\n".join(buf).strip())
            buf = [ln] if ln else []
        else:
            buf.append(ln)
    if buf:
        parts.append("\n".join(buf).strip())
    return [p for p in parts if p]

def _headers():
    return {"Authorization": f"Bearer {OPENAI_API_KEY}", "Content-Type": "application/json"}

def embed_batch(batch: List[str]) -> List[List[float]]:
    payload = {"model": EMBED_MODEL, "input": batch}
    backoff = 2
    for _ in range(5):
        r = requests.post(f"{OPENAI_BASE}/embeddings", json=payload, headers=_headers(), timeout=60)
        if r.status_code in (429, 500, 502, 503, 504):
            time.sleep(backoff); backoff = min(backoff * 2, 16); continue
        r.raise_for_status()
        return [d["embedding"] for d in r.json()["data"]]
    r.raise_for_status()

def embed_texts(texts: List[str], batch_size: int = 16) -> np.ndarray:
    vecs: List[List[float]] = []
    for i in range(0, len(texts), batch_size):
        vecs.extend(embed_batch(texts[i:i + batch_size]))
    return np.array(vecs, dtype=np.float32)

def cosine_sim_matrix(A: np.ndarray, b: np.ndarray) -> np.ndarray:
    A = A / (np.linalg.norm(A, axis=1, keepdims=True) + 1e-9)
    b = b / (np.linalg.norm(b) + 1e-9)
    return A @ b

def llm_chat(messages: List[Dict]) -> str:
    payload = {"model": CHAT_MODEL, "messages": messages, "temperature": 0.1}
    r = requests.post(f"{OPENAI_BASE}/chat/completions", json=payload, headers=_headers(), timeout=120)
    if r.status_code == 401: return "OpenAI authentication failed (401). Check your API key."
    if r.status_code == 429: return "Rate limit (429). Please retry in a moment."
    r.raise_for_status()
    return r.json()["choices"][0]["message"]["content"].strip()

# ----------------------- Helpers: read files -----------------------
def read_docx_bytes(data: bytes) -> str:
    try:
        from docx import Document
    except Exception:
        return ""
    f = io.BytesIO(data)
    doc = Document(f)
    parts: List[str] = []
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt: parts.append(txt)
    for tbl in doc.tables:
        for row in tbl.rows:
            row_txt = " | ".join([c.text.strip() for c in row.cells if c.text.strip()])
            if row_txt: parts.append(row_txt)
    return "\n".join(parts)

def read_txt_bytes(data: bytes) -> str:
    return data.decode("utf-8", errors="ignore")

def persist_kb(chunks: List[str], embeds: np.ndarray, meta: List[Dict]):
    CHUNKS_JSON.write_text(json.dumps(chunks, ensure_ascii=False, indent=2), encoding="utf-8")
    META_JSON.write_text(json.dumps(meta,   ensure_ascii=False, indent=2), encoding="utf-8")
    np.save(EMB_NPY, embeds)

def load_persisted_kb() -> bool:
    if CHUNKS_JSON.exists() and EMB_NPY.exists() and META_JSON.exists():
        try:
            chunks = json.loads(CHUNKS_JSON.read_text(encoding="utf-8"))
            meta   = json.loads(META_JSON.read_text(encoding="utf-8"))
            embeds = np.load(EMB_NPY)
            if isinstance(chunks, list) and isinstance(meta, list) and isinstance(embeds, np.ndarray):
                KB["chunks"] = chunks; KB["embeds"] = embeds; KB["meta"] = meta
                return True
        except Exception:
            return False
    return False

load_persisted_kb()

# ----------------------- API: Basic -----------------------
@app.get("/")
def home():
    return {
        "ok": True,
        "msg": "Policy Chatbot API",
        "endpoints": ["/api/health", "/api/kb/status", "/api/kb/preview",
                      "/api/kb/load (POST)", "/api/kb/clear (POST)",
                      "/api/chat (POST)", "/static/chat.html", "/docs"],
    }

@app.get("/api/health")
def health():
    return {"ok": True, "model": CHAT_MODEL, "embed_model": EMBED_MODEL,
            "kb_loaded": KB["embeds"] is not None}

@app.get("/api/kb/status")
def kb_status():
    return {
        "ok": KB["embeds"] is not None and len(KB["chunks"]) > 0,
        "chunks": len(KB["chunks"]) if KB["chunks"] else 0,
        "persisted": CHUNKS_JSON.exists() and EMB_NPY.exists() and META_JSON.exists(),
    }

@app.get("/api/kb/preview")
def kb_preview(n: int = 5):
    if KB["embeds"] is None or not KB["chunks"]:
        return {"ok": False, "msg": "KB not loaded"}
    n = max(1, min(n, len(KB["chunks"])))
    return {
        "ok": True,
        "total_chunks": len(KB["chunks"]),
        "first_sizes": [len(KB["chunks"][i]) for i in range(n)],
        "first_meta":  KB["meta"][:n],
        "first_chunks": KB["chunks"][:n],
    }

@app.post("/api/kb/clear")
def kb_clear():
    KB["chunks"] = []; KB["meta"] = []; KB["embeds"] = None
    for p in (CHUNKS_JSON, META_JSON, EMB_NPY):
        try:
            if p.exists(): p.unlink()
        except Exception:
            pass
    return {"ok": True, "msg": "KB cleared"}

# ----------------------- API: KB Load -----------------------
@app.post("/api/kb/load")
async def kb_load(
    files: List[UploadFile] = File(default=None, description="One or many .docx/.txt files"),
    text: str = Form(default=None),
    mode: Literal["append", "replace"] = Form(default="replace")
):
    """
    Load/refresh the knowledge base.
    - You may upload one or many files via 'files'
    - Or send raw 'text'
    - mode='append' adds to existing KB; 'replace' clears and rebuilds
    """
    docs: List[Dict] = []

    if text and text.strip():
        docs.append({"name": "inline_text.txt", "text": text})

    if files:
        for f in files:
            data = await f.read()
            name = (f.filename or "upload").lower()
            if name.endswith(".docx"):
                parsed = read_docx_bytes(data)
            elif name.endswith(".txt"):
                parsed = read_txt_bytes(data)
            else:
                return {"ok": False, "msg": f"Unsupported file type: {f.filename}"}
            docs.append({"name": f.filename, "text": parsed})

    if not docs:
        return {"ok": False, "msg": "Provide text or upload .docx/.txt files."}

    # Build chunks + meta per document
    new_chunks: List[str] = []
    new_meta:   List[Dict] = []
    for d in docs:
        raw = _clean_text(d["text"])
        if not raw:
            continue
        pieces = chunk_text(raw, max_chars=CHUNK_SIZE)
        for j, ch in enumerate(pieces, start=1):
            # IMPORTANT: do NOT inject [source: ...] into chunk text
            new_chunks.append(ch)
            new_meta.append({"source": d["name"], "piece": j, "total_in_doc": len(pieces)})

    if not new_chunks:
        return {"ok": False, "msg": "No usable text after parsing."}

    new_embeds = embed_texts(new_chunks)

    if mode == "append" and KB["embeds"] is not None:
        KB["chunks"].extend(new_chunks)
        KB["meta"].extend(new_meta)
        KB["embeds"] = np.vstack([KB["embeds"], new_embeds])
    else:
        KB["chunks"] = new_chunks
        KB["meta"]   = new_meta
        KB["embeds"] = new_embeds

    persist_kb(KB["chunks"], KB["embeds"], KB["meta"])
    return {"ok": True, "mode": mode, "chunks_added": len(new_chunks), "chunks_total": len(KB["chunks"])}

# ----------------------- API: Chat -----------------------
class ChatIn(BaseModel):
    message: str

SYSTEM_PROMPT = (
    "You are 'Conflict Resolution Assistant'. "
    "Answer ONLY using the policy context provided. "
    "If the answer is not clearly supported by the policy, reply: "
    "'I don't have that information in the policy.' Be concise."
)

# safety strip for any old data that still has [source: ...]
SOURCE_TAG_RE = re.compile(r'^\[source:[^\]]+\]\s*', re.IGNORECASE)

@app.post("/api/chat")
def api_chat(inp: ChatIn):
    q = (inp.message or "").strip()
    if not q:
        return {"reply": "Please enter a question."}

    if KB["embeds"] is None or len(KB["chunks"]) == 0:
        return {"reply": "The policy hasn’t been loaded yet. Please upload it to /api/kb/load."}

    try:
        qvec = embed_texts([q])[0]
    except requests.HTTPError as e:
        if e.response.status_code == 401:
            return {"reply": "OpenAI authentication failed (401). Check your API key."}
        if e.response.status_code == 429:
            return {"reply": "Rate limit (429). Please retry in a moment."}
        return {"reply": f"Embedding error: {e.response.status_code}"}
    except Exception as e:
        return {"reply": f"Embedding error: {str(e)}"}

    sims = cosine_sim_matrix(KB["embeds"], qvec)
    idx = np.argsort(-sims)[:TOP_K]

    # Build clean context (strip any legacy [source: ...] tags if they exist)
    top_chunks = []
    for i in idx:
        ch = KB["chunks"][int(i)]
        ch = SOURCE_TAG_RE.sub("", ch)
        top_chunks.append(ch)
    context = "\n\n---\n\n".join(top_chunks)
    cites   = [KB["meta"][int(i)] for i in idx]

    user_prompt = (
        "Answer the user's question strictly from the policy context below.\n\n"
        "### Policy Context\n"
        f"{context}\n\n"
        "### Question\n"
        f"{q}\n\n"
        "### Instructions\n"
        "- Quote or paraphrase the relevant rule.\n"
        "- If not in the context, say you don't have that information.\n"
        "- Keep under 120 words."
    )

    messages = [
        {"role": "system", "content": SYSTEM_PROMPT},
        {"role": "user", "content": user_prompt},
    ]
    answer = llm_chat(messages)

    # lightweight citations
    #sources = list({c['source'] for c in cites})
    #if sources:
        #answer += "\n\nSources: " + ", ".join(sources[:3])

    return {"reply": answer}
